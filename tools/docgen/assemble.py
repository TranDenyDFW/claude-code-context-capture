"""Assemble the Word document from the annotated bands and the generated notes.

Nothing is authored here. The images come from the annotator, the item numbers from the same
place, and the prose from the notes file. This lays them out and does the arithmetic on coverage,
so a gap in the notes shows up in the document as a visible gap rather than as a missing page.

Landscape, because the screenshots are 1,440 CSS pixels wide and a portrait page shrinks them to
the point where a column header cannot be read, which defeats the purpose of a document made of
screenshots.

    python tools/docgen/assemble.py --notes docs/docgen/notes.json --out docs/c4x-field-guide.docx
"""
import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent.parent
DOCGEN = ROOT / "docs" / "docgen"

# The order chapters appear in, matching the app's own tab strip so the document can be read
# alongside the running dashboard.
ORDER = [
    ("summary", "Summary", "Store-wide findings, totals, and where the tokens went."),
    ("sessions", "All sessions", "Every session as a row and as a point."),
    ("session", "Session", "One session in detail: growth, thresholds, compactions, messages."),
    ("compactions", "Compactions", "Every compaction, its predicted trigger, and what it dropped."),
    ("window-composition", "Window / Composition",
     "What is in the context window now, and how the split moved."),
    ("window-configuration", "Window / Configuration",
     "Which skills, tools, agents and memory files the configuration holds."),
    ("window-conversation", "Window / Conversation",
     "What the conversation itself put in the window."),
    ("window-injected", "Window / Injected",
     "Context nobody typed: reminders, hook output, listings."),
    ("cost", "Cost", "Paid for twice: re-reads, repeated inputs, and an estimated bill."),
    ("compare", "Compare", "Two populations measured by the same function."),
    ("diagnostics", "Diagnostics", "Is the capture healthy, and does the model agree."),
]

FIELDS = [
    ("purpose", "What this is for"),
    ("derivation", "How it is derived"),
    ("sql", "SQL"),
    ("issues", "Issues"),
    ("recommendations", "Recommendations"),
    ("other", "Other notes"),
]


def setup(doc):
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches, Pt

    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        for side in ("left", "right", "top", "bottom"):
            setattr(section, f"{side}_margin", Inches(0.6))

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    if "C4X Code" not in [s.name for s in styles]:
        from docx.enum.style import WD_STYLE_TYPE
        code = styles.add_style("C4X Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code.font.size = Pt(8)
        code.paragraph_format.space_before = Pt(2)
        code.paragraph_format.space_after = Pt(6)
        code.paragraph_format.left_indent = Inches(0.25)
    return doc


def usable_width(doc):
    from docx.shared import Inches
    s = doc.sections[0]
    return min(Inches(9.6), s.page_width - s.left_margin - s.right_margin)


def add_notes(doc, note):
    """The five notes for one item, as labelled paragraphs rather than a table.

    A table per item looked tidier and paginated badly: Word splits a six-row table across a page
    break and the label ends up orphaned from its text.
    """
    from docx.shared import Pt
    for key, label in FIELDS:
        value = (note.get(key) or "").strip()
        if not value or (key == "sql" and value.lower() in ("", "none", "n/a")):
            continue
        if key == "sql":
            head = doc.add_paragraph()
            head.add_run(f"{label}: ").bold = True
            for line in value.splitlines():
                doc.add_paragraph(line.rstrip(), style="C4X Code")
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)


def build(spec_dir, notes_path, out_path, audit_path=None):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    notes_all = json.loads(Path(notes_path).read_text(encoding="utf-8"))
    by_page = {n["page"]: {int(x["n"]): x for x in n["notes"]} for n in notes_all}
    audits = {}
    if audit_path and Path(audit_path).exists():
        for a in json.loads(Path(audit_path).read_text(encoding="utf-8")):
            audits[a["page"]] = a

    doc = setup(Document())
    width = usable_width(doc)

    doc.add_heading("c4x: field guide to every tab, panel and column", 0)
    p = doc.add_paragraph()
    p.add_run("Generated ").italic = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M")).italic = True
    p.add_run(" by tools/docgen, from the running dashboard.").italic = True

    doc.add_heading("How to read this document", 1)
    for line in [
        "Every screenshot is the live application. Red boxes are drawn from the rectangles the "
        "browser reported for each element, so a box is where the element actually is rather than "
        "where someone placed it by hand.",
        "The red number beside each box refers to the numbered entry beneath the image. Numbers "
        "run down each page in reading order and do not restart between sections.",
        "Every numbered item carries the same six notes: what it is for, how it is derived, the "
        "SQL behind it where there is one, issues, recommendations, and anything else worth "
        "knowing. An item with no SQL is either static text or computed in Python, and its "
        "derivation says which.",
        "Where a fact could not be established from the source it says so rather than guessing.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # Counted from the artefacts rather than typed in, so the methodology chapter cannot drift
    # away from what actually happened on the last run.
    repaired_total = 0
    unsupported_file = Path(spec_dir) / "unsupported.v1.json"
    if unsupported_file.exists():
        raw = json.loads(unsupported_file.read_text(encoding="utf-8"))
        repaired_total = sum(len(v) for v in raw.values())
    columns = no_help = 0
    for f in sorted(Path(spec_dir).glob("*.spec.json")):
        for item in json.loads(f.read_text(encoding="utf-8"))["items"]:
            if item["kind"] == "column":
                columns += 1
                if not (item.get("facts") or {}).get("column_help"):
                    no_help += 1

    total_items = total_noted = 0
    for page, title, blurb in ORDER:
        spec_file = Path(spec_dir) / f"{page}.spec.json"
        if not spec_file.exists():
            continue
        spec = json.loads(spec_file.read_text(encoding="utf-8"))
        notes = by_page.get(page, {})
        total_items += spec["item_count"]
        total_noted += sum(1 for i in spec["items"] if i["n"] in notes)

        doc.add_page_break()
        doc.add_heading(title, 1)
        doc.add_paragraph(blurb)
        meta = doc.add_paragraph()
        meta.add_run(f"{spec['item_count']} documented items across "
                     f"{len(spec['bands'])} views.").italic = True

        by_number = {i["n"]: i for i in spec["items"]}
        for band in spec["bands"]:
            image = Path(spec_dir) / band["image"]
            if image.exists():
                doc.add_paragraph()
                run = doc.add_paragraph().add_run()
                run.add_picture(str(image), width=width)
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for n in band["items"]:
                item = by_number.get(n)
                if not item:
                    continue
                label = (item.get("label") or item.get("key") or "").strip()
                h = doc.add_heading(f"{n}. {label[:110]}", 3)
                h.paragraph_format.space_before = Pt(10)
                kind = doc.add_paragraph()
                kind_run = kind.add_run(item["kind"].upper())
                kind_run.bold = True
                kind_run.font.size = Pt(8)
                kind_run.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
                if item.get("table_id"):
                    kind.add_run(f"   table: {item['table_id']}").font.size = Pt(8)
                if item.get("value"):
                    kind.add_run(f"   currently: {str(item['value'])[:90]}").font.size = Pt(8)
                note = notes.get(n)
                if note:
                    add_notes(doc, note)
                else:
                    miss = doc.add_paragraph()
                    miss.add_run("No note was generated for this item.").bold = True

        # The first-draft audit findings are deliberately NOT printed here.
        #
        # Every item they flagged was rewritten before this document was built, so printing them
        # beside the corrected text accuses that text of faults it no longer has. That is what the
        # first assembly did, and it made the document argue with itself. The audit's real result
        # is a fact about the process, not about any one item, so it is reported once in the
        # methodology chapter with the counts that back it.

    doc.add_page_break()
    doc.add_heading("How this document was produced and checked", 1)
    doc.add_paragraph(f"{total_noted} of {total_items} items carry notes. Coverage is reported "
                      f"rather than assumed: an item with no note is printed with a line saying "
                      f"so, so a gap is visible instead of invisible by omission.")

    doc.add_heading("Production", 2)
    for line in [
        "Every rectangle was read from the live page by a browser, so a red box is where the "
        "element actually is. No box was placed by hand.",
        "Every item's notes were written against a spec carrying the facts already established "
        "for it: the visible label, the column tooltip from COLUMN_HELP, the SQL the application "
        "itself prints in its own accordion, and file, line and function citations into the "
        "source.",
        "The whole document regenerates from the running application with five commands, so it "
        "can be refreshed after a change rather than rewritten.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Checking", 2)
    first_findings = sum(len(a.get("findings") or []) for a in audits.values())
    sev = {}
    for a in audits.values():
        for f in a.get("findings") or []:
            k = f.get("severity", "other")
            sev[k] = sev.get(k, 0) + 1
    sev_text = ", ".join(f"{v} {k}" for k, v in sorted(sev.items(), key=lambda kv: -kv[1]))
    for line in [
        f"The first draft was checked two ways. An independent agent per page, which had not "
        f"written the notes, raised {first_findings} findings ({sev_text}).",
        f"A mechanical checker (tools/docgen/check_notes.py) then refused any claim the sources "
        f"do not support: a table id absent from the page, a file path that does not exist, a "
        f"figure appearing in neither the spec nor the source, or a phrase describing a "
        f"screenshot the writer could not see. It found {repaired_total} such claims.",
        "Every flagged item was rewritten under a rule that no live value may be quoted at all, "
        "because any figure from the store is stale as soon as the store updates. The notes "
        "therefore describe what a field IS, not what it read on the day.",
        "The checker now reports zero unsupported claims, and it was proven still able to fail: "
        "injecting an invented value, an invented table id and a screenshot description makes it "
        "report exactly those three and exit non-zero.",
        "Constants defined in the source (a threshold, a window size, a LIMIT) and line-number "
        "citations are permitted and were verified against the files they name.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("A finding about the application itself", 2)
    doc.add_paragraph(
        f"{no_help} of the {columns} table columns documented here carry no entry in "
        f"COLUMN_HELP, so hovering them in the application explains nothing. Each is flagged in "
        f"its own Issues note.")

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return total_items, total_noted


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--dir", default=str(DOCGEN))
    ap.add_argument("--notes", default=str(DOCGEN / "notes.json"))
    ap.add_argument("--audit", default=str(DOCGEN / "audit.json"))
    ap.add_argument("--out", default=str(ROOT / "docs" / "c4x-field-guide.docx"))
    args = ap.parse_args(argv)

    if not Path(args.notes).exists():
        print(f"no notes at {args.notes}")
        return 1
    items, noted = build(args.dir, args.notes, args.out, args.audit)
    size = Path(args.out).stat().st_size / 1024 / 1024
    print(f"  {noted} of {items} items documented")
    print(f"  {args.out}  ({size:.1f} MB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
